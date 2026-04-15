"""
extractor.py
------------
Handles all document extraction:
  - PDF  → text + tables (pdfplumber)
  - DOCX → text (python-docx)
  - Image → OCR text (pytesseract)
"""

import io
import time
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document


# ─────────────────────────────────────────────
# PDF EXTRACTION
# ─────────────────────────────────────────────

def extract_pdf(file_bytes: bytes) -> dict:
    """
    Extract text, tables, and metadata from a PDF file.

    Args:
        file_bytes: Raw bytes of the uploaded PDF.

    Returns:
        dict with keys: text, tables, page_count, processing_time_ms
    """
    start = time.time()
    full_text = []
    all_tables = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)

        for page_num, page in enumerate(pdf.pages, start=1):
            # Extract plain text
            page_text = page.extract_text()
            if page_text:
                full_text.append(page_text.strip())

            # Extract tables → list of list of rows
            tables = page.extract_tables()
            for table in tables:
                if table:
                    cleaned_table = []
                    for row in table:
                        cleaned_row = [cell.strip() if cell else "" for cell in row]
                        cleaned_table.append(cleaned_row)
                    all_tables.append({
                        "page": page_num,
                        "data": cleaned_table
                    })

    elapsed_ms = round((time.time() - start) * 1000, 2)

    return {
        "text": "\n\n".join(full_text),
        "tables": all_tables,
        "page_count": page_count,
        "processing_time_ms": elapsed_ms
    }


# ─────────────────────────────────────────────
# DOCX EXTRACTION
# ─────────────────────────────────────────────

def extract_docx(file_bytes: bytes) -> dict:
    """
    Extract text and paragraph count from a DOCX file.

    Args:
        file_bytes: Raw bytes of the uploaded DOCX.

    Returns:
        dict with keys: text, paragraph_count, processing_time_ms
    """
    start = time.time()
    doc = Document(io.BytesIO(file_bytes))

    paragraphs = []
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            paragraphs.append(stripped)

    # Also extract text from tables inside DOCX
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                table_texts.append(row_text)

    all_text_parts = paragraphs + table_texts
    elapsed_ms = round((time.time() - start) * 1000, 2)

    return {
        "text": "\n\n".join(all_text_parts),
        "paragraph_count": len(paragraphs),
        "processing_time_ms": elapsed_ms
    }


# ─────────────────────────────────────────────
# IMAGE OCR EXTRACTION
# ─────────────────────────────────────────────

def extract_image(file_bytes: bytes) -> dict:
    """
    Perform OCR on an uploaded image and return extracted text.

    Args:
        file_bytes: Raw bytes of the uploaded image (PNG, JPG, etc.)

    Returns:
        dict with keys: text, confidence, processing_time_ms
    """
    start = time.time()
    image = Image.open(io.BytesIO(file_bytes))

    # Convert to RGB if needed (handles PNG with alpha channel)
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")

    # Run OCR — get both text and detailed data for confidence
    ocr_text = pytesseract.image_to_string(image).strip()

    # Calculate average OCR confidence
    try:
        ocr_data = pytesseract.image_to_data(
            image, output_type=pytesseract.Output.DICT
        )
        confidences = [
            int(c) for c in ocr_data["conf"] if str(c).isdigit() and int(c) > 0
        ]
        avg_confidence = round(sum(confidences) / len(confidences) / 100, 2) if confidences else 0.0
    except Exception:
        avg_confidence = 0.0

    elapsed_ms = round((time.time() - start) * 1000, 2)

    return {
        "text": ocr_text,
        "confidence": avg_confidence,   # 0.0 – 1.0
        "processing_time_ms": elapsed_ms
    }


# ─────────────────────────────────────────────
# UNIFIED ROUTER
# ─────────────────────────────────────────────

def extract(file_bytes: bytes, file_type: str) -> dict:
    """
    Route extraction based on file type.

    Args:
        file_bytes : Raw bytes of the file.
        file_type  : One of 'pdf', 'docx', 'image'

    Returns:
        Extraction result dict (varies by type).

    Raises:
        ValueError: If file_type is not supported.
    """
    file_type = file_type.lower().strip()

    if file_type == "pdf":
        return extract_pdf(file_bytes)
    elif file_type in ("docx", "doc"):
        return extract_docx(file_bytes)
    elif file_type in ("image", "png", "jpg", "jpeg", "bmp", "tiff", "webp"):
        return extract_image(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type: '{file_type}'. "
            "Supported types: pdf, docx, image (png/jpg/jpeg/bmp/tiff/webp)"
        )
